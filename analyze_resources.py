import docx
import re
import json
import os

def extract_unit_resources(doc_path):
    """Extract resources from DOCX file by unit."""
    # Check if file exists
    if not os.path.exists(doc_path):
        print(f"File not found: {doc_path}")
        return None
    
    try:
        # Open the document
        doc = docx.Document(doc_path)
        
        # Extract text
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Find unit sections
        unit_pattern = r'VISUAL 1 - UNIT (\d+) - (.+?)(?=VISUAL 1 - UNIT \d+|\Z)'
        matches = re.findall(unit_pattern, full_text, re.DOTALL)
        
        if not matches:
            print("No units found in the document")
            # Try another approach - extract paragraphs that have 'UNIT' in them
            unit_texts = []
            current_unit = ""
            current_unit_num = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if "VISUAL 1 - UNIT" in text:
                    # If we were building a unit, save it
                    if current_unit_num:
                        unit_texts.append((current_unit_num, current_unit))
                    
                    # Start a new unit
                    unit_match = re.search(r'UNIT (\d+)', text)
                    if unit_match:
                        current_unit_num = unit_match.group(1)
                        current_unit = text + "\n"
                    else:
                        current_unit_num = None
                        current_unit = ""
                elif current_unit_num:
                    current_unit += text + "\n"
            
            # Add the last unit
            if current_unit_num:
                unit_texts.append((current_unit_num, current_unit))
            
            matches = unit_texts
        
        # Process each unit
        resource_data = {}
        for unit_num, unit_content in matches:
            # Extract title
            title_match = re.search(r'UNIT \d+ - ([^\n]+)', unit_content)
            title = title_match.group(1).strip() if title_match else "Unknown"
            
            # Extract videos
            video_patterns = [
                r'VIDEO ([^\n<]+?)(?:\n|$)',
                r'VIDEO SONG ([^\n<]+?)(?:\n|$)',
                r'VIDEO FILM ([^\n<]+?)(?:\n|$)',
                r'VIDEO SKIT ([^\n<]+?)(?:\n|$)'
            ]
            
            videos = []
            for pattern in video_patterns:
                videos.extend(re.findall(pattern, unit_content))
            
            # Extract games
            game_pattern = r'ONLINE GAME ([^\n<]+?)(?:\n|$)'
            games = re.findall(game_pattern, unit_content)
            
            # Count iframes for each resource type
            iframe_pattern = r'<iframe.+?</iframe>'
            iframes = re.findall(iframe_pattern, unit_content)
            
            # Count YouTube and Wordwall iframes
            youtube_count = sum(1 for iframe in iframes if 'youtube.com/embed' in iframe)
            wordwall_count = sum(1 for iframe in iframes if 'wordwall.net/embed' in iframe)
            
            # Store data for this unit
            resource_data[unit_num] = {
                'title': title,
                'videos': videos,
                'games': games,
                'video_count': len(videos),
                'game_count': len(games),
                'iframe_count': len(iframes),
                'youtube_count': youtube_count,
                'wordwall_count': wordwall_count
            }
        
        return resource_data
    
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return None

def print_resource_table(resource_data):
    """Print a formatted table of resource data."""
    if not resource_data:
        print("No resource data available")
        return
    
    print("\nVISUAL ENGLISH BOOK 1 RESOURCE INVENTORY:")
    print("=" * 78)
    print(f"{'UNIT':<6}{'TITLE':<30}{'VIDEOS':<8}{'GAMES':<8}{'YOUTUBE':<9}{'WORDWALL':<10}")
    print("-" * 78)
    
    for unit_num in sorted(resource_data.keys(), key=lambda x: int(x) if x.isdigit() else 999):
        unit = resource_data[unit_num]
        print(f"{unit_num:<6}{unit['title'][:28]:<30}{unit['video_count']:<8}{unit['game_count']:<8}" 
              f"{unit['youtube_count']:<9}{unit['wordwall_count']:<10}")
    
    print("=" * 78)
    
    # Print detailed resources for generation
    print("\nDETAILED RESOURCES FOR CODE GENERATION:")
    print("=" * 100)
    
    for unit_num in sorted(resource_data.keys(), key=lambda x: int(x) if x.isdigit() else 999):
        unit = resource_data[unit_num]
        print(f"\nUNIT {unit_num}:")
        
        # Get full text for this unit from docx file
        with open('downloaded_document.docx', 'rb') as docx_file:
            doc = docx.Document(docx_file)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            unit_pattern = f'VISUAL 1 - UNIT {unit_num} - (.+?)(?=VISUAL 1 - UNIT \d+|\Z)'
            unit_match = re.search(unit_pattern, full_text, re.DOTALL)
            unit_content = unit_match.group(0) if unit_match else ""
        
        # Extract and print all iframe codes
        iframe_pattern = r'<iframe.+?</iframe>'
        iframes = re.findall(iframe_pattern, unit_content)
        
        print("  Videos:")
        for i, iframe in enumerate([iframe for iframe in iframes if 'youtube.com/embed' in iframe]):
            # Extract video ID
            video_id_match = re.search(r'youtube.com/embed/([^?]+)', iframe)
            video_id = video_id_match.group(1) if video_id_match else "unknown"
            
            # Attempt to match with a video title
            title = f"Video {i+1}"
            if i < len(unit['videos']):
                title = unit['videos'][i]
            
            print(f"    {i+1}. {title}")
            print(f"       URL: https://www.youtube.com/watch?v={video_id}")
            print(f"       Embed: {iframe}")
        
        print("\n  Games:")
        for i, iframe in enumerate([iframe for iframe in iframes if 'wordwall.net/embed' in iframe]):
            # Extract game ID
            game_id_match = re.search(r'wordwall.net/embed/([^?]+)', iframe)
            game_id = game_id_match.group(1) if game_id_match else "unknown"
            
            # Attempt to match with a game title
            title = f"Game {i+1}"
            if i < len(unit['games']):
                title = unit['games'][i]
            
            print(f"    {i+1}. {title}")
            print(f"       URL: https://wordwall.net/resource/{game_id}")
            print(f"       Embed: {iframe}")
            
    print("\n" + "=" * 100)

def compare_with_implementation(resource_data):
    """Compare source document counts with implementation counts."""
    # Hardcoded implementation counts based on our earlier analysis
    implementation = {
        '1': {'videos': 3, 'games': 5},
        '2': {'videos': 6, 'games': 3},
        '3': {'videos': 5, 'games': 2},
        '4': {'videos': 4, 'games': 2},
        '5': {'videos': 5, 'games': 3},
        '6': {'videos': 4, 'games': 2},  # Updated after our fix
        '7': {'videos': 3, 'games': 1},
        '8': {'videos': 3, 'games': 1},
        '9': {'videos': 3, 'games': 1},
        '10': {'videos': 3, 'games': 2},
        '11': {'videos': 3, 'games': 2},
        '12': {'videos': 3, 'games': 2},
        '13': {'videos': 3, 'games': 2},
        '14': {'videos': 3, 'games': 2},
        '15': {'videos': 3, 'games': 2},
        '16': {'videos': 3, 'games': 2},
        '17': {'videos': 3, 'games': 2},
        '18': {'videos': 3, 'games': 2}
    }
    
    print("\nCOMPARISON: SOURCE DOCUMENT vs IMPLEMENTATION")
    print("=" * 95)
    print(f"{'UNIT':<6}{'TITLE':<25}{'SRC VIDEOS':<12}{'IMPL VIDEOS':<13}{'SRC GAMES':<11}{'IMPL GAMES':<12}{'STATUS':<10}")
    print("-" * 95)
    
    for unit_num in sorted(resource_data.keys(), key=lambda x: int(x) if x.isdigit() else 999):
        if unit_num not in implementation:
            continue
            
        unit = resource_data[unit_num]
        src_videos = unit['youtube_count']
        src_games = unit['wordwall_count']
        impl_videos = implementation[unit_num]['videos']
        impl_games = implementation[unit_num]['games']
        
        # Determine status
        if src_videos == impl_videos and src_games == impl_games:
            status = "Complete"
        elif src_videos == impl_videos or src_games == impl_games:
            status = "Partial"
        else:
            status = "Incomplete"
            
        print(f"{unit_num:<6}{unit['title'][:23]:<25}{src_videos:<12}{impl_videos:<13}{src_games:<11}{impl_games:<12}{status:<10}")
    
    print("=" * 95)

def main():
    doc_path = 'downloaded_document.docx'
    resource_data = extract_unit_resources(doc_path)
    
    if resource_data:
        print_resource_table(resource_data)
        compare_with_implementation(resource_data)
        
        # Save detailed data as JSON
        with open('resource_inventory.json', 'w') as json_file:
            json.dump(resource_data, json_file, indent=2)
            print("\nDetailed resource inventory saved to resource_inventory.json")

if __name__ == "__main__":
    main()