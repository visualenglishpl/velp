import docx
import re
import json
import os

def extract_iframes(docx_path):
    """Extract all iframes from a DOCX file, organized by unit."""
    try:
        # Open the document
        doc = docx.Document(docx_path)
        
        # Extract all paragraphs with iframes
        iframes_by_unit = {}
        current_unit = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Check if this is a unit header
            unit_match = re.search(r'VISUAL 1 - UNIT (\d+)', text)
            if unit_match:
                current_unit = unit_match.group(1)
                if current_unit not in iframes_by_unit:
                    iframes_by_unit[current_unit] = {
                        'videos': [],
                        'games': [],
                        'title': text.replace('VISUAL 1 - UNIT ' + current_unit + ' - ', '').strip()
                    }
            
            # Check if this is an iframe
            iframe_match = re.search(r'<iframe.+?</iframe>', text)
            if iframe_match and current_unit:
                iframe = iframe_match.group(0)
                
                # Determine if it's a video or game
                if 'youtube.com/embed' in iframe:
                    # Extract title from previous paragraph if available
                    title = "Video"
                    video_id_match = re.search(r'youtube.com/embed/([^?"]+)', iframe)
                    video_id = video_id_match.group(1) if video_id_match else "unknown"
                    
                    # Look for title in previous paragraph
                    idx = doc.paragraphs.index(para)
                    if idx > 0:
                        prev_text = doc.paragraphs[idx-1].text.strip()
                        # If previous text has VIDEO in it, it's likely the title
                        if re.search(r'VIDEO|SONG|FILM|SKIT', prev_text):
                            title = prev_text
                    
                    iframes_by_unit[current_unit]['videos'].append({
                        'title': title,
                        'id': video_id,
                        'url': f"https://www.youtube.com/watch?v={video_id}",
                        'embed': iframe
                    })
                elif 'wordwall.net/embed' in iframe:
                    # Extract title from previous paragraph if available
                    title = "Game"
                    game_id_match = re.search(r'wordwall.net/embed/([^?"]+)', iframe)
                    game_id = game_id_match.group(1) if game_id_match else "unknown"
                    
                    # Look for title in previous paragraph
                    idx = doc.paragraphs.index(para)
                    if idx > 0:
                        prev_text = doc.paragraphs[idx-1].text.strip()
                        # If previous text has GAME or WORDWALL in it, it's likely the title
                        if re.search(r'GAME|WORDWALL', prev_text):
                            title = prev_text
                    
                    iframes_by_unit[current_unit]['games'].append({
                        'title': title,
                        'id': game_id,
                        'url': f"https://wordwall.net/resource/{game_id}",
                        'embed': iframe
                    })
        
        return iframes_by_unit
    
    except Exception as e:
        print(f"Error extracting iframes: {str(e)}")
        return None

def write_unit_resource_file(unit_num, resources, output_dir="./unit_resources"):
    """Write extracted resources to a file for a specific unit."""
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = os.path.join(output_dir, f"unit{unit_num}_resources.json")
    
    with open(output_path, 'w') as outfile:
        json.dump(resources, outfile, indent=2)
    
    print(f"Wrote resources for Unit {unit_num} to {output_path}")

def generate_resource_files(resources_by_unit, output_dir="./unit_resources"):
    """Generate resource files for all units."""
    for unit_num, resources in resources_by_unit.items():
        write_unit_resource_file(unit_num, resources, output_dir)
    
    summary = {unit: {
        'title': data['title'],
        'video_count': len(data['videos']),
        'game_count': len(data['games'])
    } for unit, data in resources_by_unit.items()}
    
    # Write summary file
    with open(os.path.join(output_dir, "resource_summary.json"), 'w') as outfile:
        json.dump(summary, outfile, indent=2)
    
    print(f"Wrote summary to {os.path.join(output_dir, 'resource_summary.json')}")

def main():
    doc_path = 'downloaded_document.docx'
    
    if not os.path.exists(doc_path):
        print(f"Document not found: {doc_path}")
        return
    
    resources_by_unit = extract_iframes(doc_path)
    
    if resources_by_unit:
        generate_resource_files(resources_by_unit)
        
        # Print summary
        print("\nResource Summary:")
        print("=" * 60)
        print(f"{'Unit':<6}{'Title':<30}{'Videos':<8}{'Games':<8}")
        print("-" * 60)
        
        for unit_num in sorted(resources_by_unit.keys(), key=lambda x: int(x) if x.isdigit() else 999):
            resources = resources_by_unit[unit_num]
            print(f"{unit_num:<6}{resources['title'][:28]:<30}{len(resources['videos']):<8}{len(resources['games']):<8}")
        
        print("=" * 60)
    else:
        print("Failed to extract resources")

if __name__ == "__main__":
    main()